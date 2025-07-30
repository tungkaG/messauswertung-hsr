import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from asammdf import MDF
from docx import Document
from docx.shared import Inches, RGBColor
import io
import tkinter as tk
from tkinter import filedialog, messagebox, Listbox
import pyproj
import utm

# todo: check what happens when we do doc.add... a None

def _get_time_range(signals: dict) -> np.ndarray:
    # Function to determine the start and end time of the measurement
    if signals is None:
        return np.empty(shape=(0,))
    timeStart = 1e10
    timeEnd = 0
    for _, signal in signals.items():
        if signal.timestamps is None:
            continue
        try:
            timeEnd = max(timeEnd, signal.timestamps[-1])
            timeStart = min(timeStart, signal.timestamps[0])
        except IndexError as ie:
            print(ie)
    timeVector = np.arange(round(timeStart, 3), round(timeEnd, 3), 0.01)
    return timeVector

def _convert_signals(signals: dict, convert: bool):
    # Function to convert ASAM MDF signals to a unified time vector
    time_vector = _get_time_range(signals)
    if convert:
        df = pd.DataFrame(time_vector, columns=['time'])
        for key, signal in signals.items():
            if len(signal.samples) == 0:
                continue
            syncSignal = signal.interp(time_vector, 0, 0)
            dataSignal = list(zip(syncSignal.timestamps, syncSignal.samples))
            dfSignal = pd.DataFrame(dataSignal, columns=['time', key])
            df = pd.merge(df, dfSignal, on='time')
        return df
    else:
        for key, signal in signals.items():
            syncSignal = signal.interp(time_vector, 0, 0)
            signal.samples = syncSignal.samples
        signals['time'] = time_vector
        return signals

def extract_data(filename, convert=False):
    # Function to process the MDF4 file and extract the required signals
    try:
        measurement = MDF(filename)
    except Exception as e:
        print(f"Error opening the file {filename}: {e}")
        return None
    
    signalsInMeas = measurement.channels_db
    
    analysis_channel_list = [
        'Rate_Hor_Z', 
        'AVL_STEA_FTAX_WHL',
        'AVL_STEA_DV',
        'INS_Vel_Hor_X',
        'EgoIn5ms_Ins_vx_Act',
        'MobsObs_vx_Act',
        'DcrInEgoM_psid_Act',
        'DcrInEgoM_agwFA_Ste',
        'GNSSPositionDegreeOfLatitude',
        'GNSSPositionDegreeOfLongitude',
        'INS_Lat_Abs_POI1',
        'INS_Long_Abs_POI1',
        'INS_Lat_Abs',
        'INS_Long_Abs',
        'DcrInEgoM_beta_Act',
        'Ins_beta_Act',
        'QU_FN_FDR',
    ]
    
    mappedSignals = {}
    for signal in analysis_channel_list:
        for sub, indexgroup in signalsInMeas.items():
            if sub == signal:
                indexFindings = indexgroup[0]
                group = indexFindings[0]
                index = indexFindings[1]
                mappedSignals[signal] = measurement.get(signal,
                                                        group=group,
                                                        index=index)
    analysis_data_converted = _convert_signals(mappedSignals, convert)
    
    manipulated_signals_channel_list = [
        # rDyn
        'EgoWhlRtab_rFL_Whl',
        'EgoWhlRtab_rFR_Whl',
        'EgoWhlRtab_rRL_Whl',
        'EgoWhlRtab_rRR_Whl',
        'EgoWhlRtab_rFL_WhlLut',
        'EgoWhlRtab_rFR_WhlLut',
        'EgoWhlRtab_rRL_WhlLut',
        'EgoWhlRtab_rRR_WhlLut',

        # vx and vy
        # # nicht manipulierbare Signale
        'EgoIn5ms_Ins_vx_Act', 
        'EgoIn5ms_Ins_vy_Act',
        'MobsObs_vx_Act',
        'MobsObs_vy_Act',
        # # manipulierbare Signale
        # 'DcrInEgoM_v_Act',
        'DcrInEgoM_vx_Act',
        'DcrInEgoM_vy_Act', 

        # Steering angle
        'EgoSteaProc_agFA_Whl', # manipulierbare Signale
        'EgoStea_agFA_WhlRaw', # nicht manipulierbare Signale
        # 'EgoStea_agFA_WhlOfs', nutzlos
        # 'EgoStea_agFA_WhlOfsLt', nutzlos
        # 'EgoStea_agFA_WhlOfsPni', nutzlos
        # 'EgoStea_agFA_WhlOfsPsid', für gierratemanipulation

        # Slip angle
        # # Manipulierbare Signale
        'SnsMobsAdj_alphaFL_Act',
        'SnsMobsAdj_alphaFR_Act',
        'SnsMobsAdj_alphaRL_Act',
        'SnsMobsAdj_alphaRR_Act',
        # Nicht manipulierbare Signale
        'EgoOut_alpha_Act_rc.EgoOut_alphaFL_Act',
        'EgoOut_alpha_Act_rc.EgoOut_alphaFR_Act',
        'EgoOut_alpha_Act_rc.EgoOut_alphaRL_Act',
        'EgoOut_alpha_Act_rc.EgoOut_alphaRR_Act',

        # ax, ay, psid
        # nicht Manipulierbare SIganle
        "EgoIn5ms_psid_Raw1",
        "EgoIn5ms_ax_Raw1",
        "EgoIn5ms_ay_Raw1",
        # Manipulierbare Signale
        "EgoInWr_psid_Raw1",
        "EgoInWr_ax_Raw1",
        "EgoInWr_ay_Raw1",

        # mue split
        "DmcIndMs_b_IndMs",
        "DmcIndMs_b_IndMsSusp", 
        "DmcIndMs_b_IndMsSusp_Set",
        "DmcIndMs_fac_IndMs",
        "DmcIndMs_fac_IndMsSusp",
        "DmcIndMs_cw_IndMs",
        "DmcIndMs_cw_IndMs_Side",

        # Reibwert
        "SnsRc_mue_ActHi",

        # ay no grav
        "DcrInEgoM_ay_ActNoGrav",

        # TEE
        "DcrTql_tqwFA_EmMaxDyn",
        "DcrTql_tqwFA_EmMaxStat",
        "DcrTql_tqwFA_EmMinDyn",
        "DcrTql_tqwFA_EmMinStat",
        "DcrTql_tqwRA_EmMaxDyn",
        "DcrTql_tqwRA_EmMaxStat",
        "DcrTql_tqwRA_EmMinDyn",
        "DcrTql_tqwRA_EmMinStat",

        # # vch
        # "EgoOutWr_rc_v_Ch.EgoOut_v_Ch",
        # "DcrInEgoE_vch_Veh",
    ]

    mappedSignals = {}
    for signal in manipulated_signals_channel_list:
        for sub, indexgroup in signalsInMeas.items():
            if sub == signal:
                indexFindings = indexgroup[0]
                group = indexFindings[0]
                index = indexFindings[1]
                mappedSignals[signal] = measurement.get(signal,
                                                        group=group,
                                                        index=index)
    manipulated_signals_data_converted = _convert_signals(mappedSignals, convert)

    print("Processing analysis data from file:", filename)
    return analysis_data_converted, manipulated_signals_data_converted

def process_mf4_analysis_data(mf4_data, filename, progress_listbox):
    warnings = []
    plot_buf = []

    # Function to process the MF4 data and perform analysis
    if mf4_data is None or mf4_data.empty:
        progress_listbox.insert(tk.END, f"MF4 data is empty or invalid for file: {filename}")
        return None

    # Find signal for psid
    signal_name_psi_d = ""
    unit_psi_d = ""
    if "Rate_Hor_Z" in mf4_data.columns:
        signal_name_psi_d = "Rate_Hor_Z"
        unit_psi_d = "°/s"
    elif "DcrInEgoM_psid_Act" in mf4_data.columns:
        progress_listbox.insert(tk.END, f"Rate_Hor_Z not available for file: {filename}, using DcrInEgoM_psid_Act instead")
        signal_name_psi_d = "DcrInEgoM_psid_Act"
        unit_psi_d = "rad/s"
    else:
        progress_listbox.insert(tk.END, f"WARNING: Rate_Hor_Z and DcrInEgoM_psid_Act not available for file: {filename}")
        warnings.append(f"WARNING: Rate_Hor_Z and DcrInEgoM_psid_Act not available for this file")
        return {
            "filename": filename,
            "plot": None,
            "psid_peak": None,
            "signal_info": None,
            "within_35_percent": None,
            "within_20_percent": None,
            "vx_begin_kmh": None,
            "warning": warnings,
            "test_passed": False
            }

    # Find signal for lenkwinkel
    signal_name_lenkwinkel = ""
    unit_lenkwinkel = ""
    if "DcrInEgoM_agwFA_Ste" in mf4_data.columns:
        signal_name_lenkwinkel = "DcrInEgoM_agwFA_Ste"
        unit_lenkwinkel = "rad"
    elif "AVL_STEA_FTAX_WHL" in mf4_data.columns:
        progress_listbox.insert(tk.END, f"DcrInEgoM_agwFA_Ste not available for file: {filename}, using AVL_STEA_FTAX_WHL instead")
        signal_name_lenkwinkel = "AVL_STEA_FTAX_WHL"
        unit_lenkwinkel = "°"
    elif "AVL_STEA_DV" in mf4_data.columns:
        progress_listbox.insert(tk.END, f"DcrInEgoM_agwFA_Ste and AVL_STEA_FTAX_WHL not available for file: {filename}, using AVL_STEA_DV instead")
        signal_name_lenkwinkel = "AVL_STEA_DV"
        unit_lenkwinkel = "°"
    else:
        progress_listbox.insert(tk.END, f"WARNING: DcrInEgoM_agwFA_Ste, AVL_STEA_FTAX_WHL, and AVL_STEA_DV not available for file: {filename}")
        warnings.append(f"WARNING: DcrInEgoM_agwFA_Ste, AVL_STEA_FTAX_WHL, and AVL_STEA_DV not available for this file")
        return {
            "filename": filename,
            "plot": None,
            "psid_peak": None,
            "signal_info": None,
            "within_35_percent": None,
            "within_20_percent": None,
            "vx_begin_kmh": None,
            "warning": warnings,
            "test_passed": False
            }

    # Find signal for vx
    signal_name_vx = ""
    unit_vx = ""
    if "INS_Vel_Hor_X" in mf4_data.columns:
        signal_name_vx = "INS_Vel_Hor_X"
        unit_vx = "m/s"
    elif "EgoIn5ms_Ins_vx_Act" in mf4_data.columns:
        progress_listbox.insert(tk.END, f"INS_Vel_Hor_X not available for file: {filename}, using EgoIn5ms_Ins_vx_Act instead")
        signal_name_vx = "EgoIn5ms_Ins_vx_Act"
        unit_vx = "m/s"
    elif "MobsObs_vx_Act" in mf4_data.columns:
        progress_listbox.insert(tk.END, f"INS_Vel_Hor_X not available for file: {filename}, using MobsObs_vx_Act instead")
        signal_name_vx = "MobsObs_vx_Act"
        unit_vx = "m/s"
    # elif "EgoMobs_Mobs_vx_Act" in mf4_data.columns:
    #     progress_listbox.insert(tk.END, f"INS_Vel_Hor_X not available for file: {filename}, using EgoMobs_Mobs_vx_Act instead")
    #     signal_name_vx = "EgoMobs_Mobs_vx_Act"
    #     unit_vx = "m/s"
    else:
        progress_listbox.insert(tk.END, f"WARNING: INS_Vel_Hor_X, EgoIn5ms_Ins_vx_Act, and MobsObs_vx_Act not available for file: {filename}")
        warnings.append(f"WARNING: INS_Vel_Hor_X, EgoIn5ms_Ins_vx_Act, and MobsObs_vx_Act not available for file: {filename}")

    # Find start of SWD (QN_FN_FDR = 512 and gradient of d_lenkwinkel/dt > 0.009)
    start_index = 0
    time_diff = round(mf4_data["time"][1] - mf4_data["time"][0],3) # 2 decs behind 0
    distance_in_index_after_0p1_s = int(0.1 / time_diff)
    for i in range(len(mf4_data[signal_name_lenkwinkel]) - distance_in_index_after_0p1_s): # search +-80 samples when QN_FN_FDR is activated, and where gradient is > 0.009
        if any(mf4_data["QU_FN_FDR"][max(i-80,0):min(i+80,len(mf4_data[signal_name_lenkwinkel]) - distance_in_index_after_0p1_s)]  != 512) and abs(mf4_data[signal_name_lenkwinkel].iloc[i + distance_in_index_after_0p1_s] - mf4_data[signal_name_lenkwinkel].iloc[i]) / 0.1 > 0.44:
            start_index = i
            break
    if start_index == 0:
        progress_listbox.insert(tk.END, f"WARNING: Could not find the start of SWD for file: {filename}")
        warnings.append(f"WARNING: Could not find the start of SWD for this file")
        return {
            "filename": filename,
            "plot": None,
            "psid_peak": None,
            "signal_info": None,
            "within_35_percent": None,
            "within_20_percent": None,
            "vx_begin_kmh": None,
            "warning": warnings,
            "test_passed": False
        }
    mf4_data_reduced = mf4_data.iloc[start_index:].reset_index(drop=True) # Trim data

    # Find maximum psid
    psid_peak_index = mf4_data_reduced[signal_name_psi_d].abs().idxmax()
    psid_peak = mf4_data_reduced[signal_name_psi_d].iloc[psid_peak_index]

    # Find time when sine sweep stopped (compare now value to last value), starting from time where we found psid_max
    T_0 = None
    state = 0
    for i in range(psid_peak_index, len(mf4_data_reduced[signal_name_lenkwinkel]) - distance_in_index_after_0p1_s):
        if state == 0:
            if abs((mf4_data_reduced[signal_name_lenkwinkel].iloc[i] - mf4_data_reduced[signal_name_lenkwinkel].iloc[i+distance_in_index_after_0p1_s])/0.1) <= 0:
                state = 1
        elif state == 1:
            if abs((mf4_data_reduced[signal_name_lenkwinkel].iloc[i] - mf4_data_reduced[signal_name_lenkwinkel].iloc[i+distance_in_index_after_0p1_s])/0.1) > 0.2:
                state = 2
        elif state == 2:
            if abs((mf4_data_reduced[signal_name_lenkwinkel].iloc[i] - mf4_data_reduced[signal_name_lenkwinkel].iloc[i+distance_in_index_after_0p1_s])/0.1) < 0.2:
                T_0 = mf4_data_reduced["time"].iloc[i]
                break

    if T_0 is None:
        progress_listbox.insert(tk.END, f"T_0 could not be determined for file: {filename}")
        warnings.append(f"T_0 could not be determined for this file")
        return {
            "filename": filename,
            "plot": None,
            "psid_peak": None,
            "signal_info": None,
            "within_35_percent": None,
            "within_20_percent": None,
            "vx_begin_kmh": None,
            "warning": warnings,
            "test_passed": False
        }

    tStartSine = mf4_data_reduced["time"].iloc[0]
    T_0 = T_0 - tStartSine

    # Create plots for psid 
    fig_psid, ax_psid = plt.subplots(figsize=(10, 6))
    ax_psid.plot(mf4_data_reduced["time"] - tStartSine, mf4_data_reduced[signal_name_psi_d], label=signal_name_psi_d + f" ({unit_psi_d})")
    ax_psid.axvline(x=T_0, color='red', linestyle='--', label=f'T_0={T_0:.3f}s')
    # plot psid peak
    psid_peak_time = mf4_data_reduced["time"].iloc[psid_peak_index] - tStartSine
    ax_psid.axvline(x=psid_peak_time, color='green', linestyle='--', label=f'Psid Peak (Time={psid_peak_time:.3f}s, Val={psid_peak:.3f} {unit_psi_d})')
    # Criterium 1 for psid
    psid_at_t0_plus_1 = mf4_data_reduced[signal_name_psi_d].iloc[
        (mf4_data_reduced["time"] - tStartSine).sub(T_0 + 1).abs().idxmin()
    ]
    ax_psid.axvline(x=T_0+1, color='blue', linestyle='--', label=f'Psid(T_0+1)={psid_at_t0_plus_1:.3f} {unit_psi_d}')
    ax_psid.fill_betweenx(
        [-psid_peak * 0.35, psid_peak * 0.35],
        T_0 + 0.95,
        T_0 + 1.05,
        color='green',
        alpha=0.3,
        label='35% Psid Peak at T_0+1'
    )
    within_35_percent = (abs(psid_at_t0_plus_1) <= abs(psid_peak) * 0.35)
    # Criteirum 2 for psid
    psid_at_t0_plus_1p75 = mf4_data_reduced[signal_name_psi_d].iloc[
        (mf4_data_reduced["time"] - tStartSine).sub(T_0 + 1.75).abs().idxmin()
    ]
    ax_psid.axvline(x=T_0+1.75, color='black', linestyle='--', label=f'Psid(T_0+1.75)={psid_at_t0_plus_1p75:.3f} {unit_psi_d}')
    ax_psid.fill_betweenx(
        [-psid_peak * 0.2, psid_peak * 0.2],
        T_0 + 1.7,
        T_0 + 1.8,
        color='green',
        alpha=0.6,
        label='20% Psid Peak at T_0+1.75'
    )
    within_20_percent = (abs(psid_at_t0_plus_1p75) <= abs(psid_peak) * 0.2)
    ax_psid.set_ylabel(signal_name_psi_d + f" ({unit_psi_d})") # label and legend for psid plot
    ax_psid.legend()
    ax_psid.grid()
    psid_buf = io.BytesIO() # save to buf
    fig_psid.savefig(psid_buf, format='png')
    plt.close(fig_psid)
    plot_buf.append(psid_buf)

    # Plot lenkwinkel
    fig_lw, ax_lw = plt.subplots(figsize=(10, 6))
    ax_lw.plot(mf4_data_reduced["time"] - tStartSine, mf4_data_reduced[signal_name_lenkwinkel], label=signal_name_lenkwinkel + f" ({unit_lenkwinkel})")
    ax_lw.axvline(x=T_0, color='red', linestyle='--')
    ax_lw.set_ylabel(signal_name_lenkwinkel + f" ({unit_lenkwinkel})")
    ax_lw.set_xlabel("Time (s)")
    ax_lw.legend()
    ax_lw.grid()
    lenkwinkel_buf = io.BytesIO()
    fig_lw.savefig(lenkwinkel_buf, format='png')
    plt.close(fig_lw)
    plot_buf.append(lenkwinkel_buf)

    # Check Querversatz (1.07s after start)
    signal_names_querversatz = [('GNSSPositionDegreeOfLatitude', 'GNSSPositionDegreeOfLongitude'), ('INS_Lat_Abs_POI1', 'INS_Long_Abs_POI1'), ('INS_Lat_Abs', 'INS_Long_Abs')]
    signal_name_querversatz = ""
    unit_querversatz = "m"

    fig_querversatz, ax_querversatz = plt.subplots(figsize=(10, 6))
    querversatz_buf = io.BytesIO()
    
    for signal_name_pair in signal_names_querversatz:
        lat_signal, lon_signal = signal_name_pair
        if lat_signal in mf4_data_reduced.columns and lon_signal in mf4_data_reduced.columns:
        
            # alle in grad
            lat1 = mf4_data_reduced[lat_signal][0]
            lon1 = mf4_data_reduced[lon_signal][0]
            lat2 = mf4_data_reduced[lat_signal]
            lon2 = mf4_data_reduced[lon_signal]

            if any(abs(lat2) > 90) or any(abs(lon2) > 180):
                warnings.append(f"WARNING: {lat_signal},{lon_signal} have unvalid values")
                continue

            signal_name_querversatz += f"{lat_signal},{lon_signal},"
            
            ########################## IN HOUSE HAVERSINE ########################
            # lat1 = np.deg2rad(lat1)
            # lon1 = np.deg2rad(lon1)
            # lat2 = np.deg2rad(lat2)
            # lon2 = np.deg2rad(lon2)

            # a = np.sin((lat2 - lat1) / 2) ** 2 + np.cos(lat1) * np.cos(lat2) * np.sin((lon2 - lon1) / 2) ** 2
            # dist = 6378.388 * 2.0 * np.arctan2(np.sqrt(a), np.sqrt(1.0-a)) * 1000 # in m
            
            # num = np.sin(lon2 - lon1) * np.cos(lat2)
            # den = np.cos(lat1) * np.sin(lat2) - np.sin(lat1) * np.cos(lat2) * np.cos(lon2 - lon1)
            # bearing = np.arctan2(num, den)

            # x = dist * np.sin(bearing)
            # y = dist * np.cos(bearing)
            ########################## IN HOUSE HAVERSINE ########################

            ########################## pyproj ####################################
            wgs84 = pyproj.CRS('EPSG:4326')
            # Automatically determine UTM zone using the first coordinate
            utm_zone = utm.from_latlon(lat1, lon1)
            utm_crs = pyproj.CRS(f'EPSG:{32600 + utm_zone[2]}')

            # Create a transformer from WGS84 to the determined UTM zone
            transformer = pyproj.Transformer.from_crs(wgs84, utm_crs)

            # Convert the array of positions to UTM
            easthing, northing = transformer.transform(lat2, lon2)
            y = easthing - easthing[0]
            x = northing - northing[0]
            ########################## pyproj ####################################

            initial_heading = np.arctan2(y[1] - y[0], x[1] - x[0])
            lateral = (x - x[0]) * np.sin(initial_heading) - (y - y[0]) * np.cos(initial_heading)
            longitudinal = (x - x[0]) * np.cos(initial_heading) + (y - y[0]) * np.sin(initial_heading)
            index_lenkwinkel_1p07_after = np.argmin(np.abs(mf4_data_reduced["time"] - (tStartSine + 1.07)))
            lateral_1p07 = lateral[index_lenkwinkel_1p07_after]

            ax_querversatz.plot(mf4_data_reduced["time"] - tStartSine, lateral)
            ax_querversatz.scatter(mf4_data_reduced["time"].iloc[index_lenkwinkel_1p07_after] - tStartSine, lateral_1p07, marker='x', color='red', label=f"Lateral Displacement at 1.07s: {lateral_1p07:.2f} {unit_querversatz}")
            ax_querversatz.set_xlabel("Time (s)")
            ax_querversatz.set_ylabel("Querversatz from " + lat_signal + "," + lon_signal + f" ({unit_querversatz})")
            ax_querversatz.grid()
            ax_querversatz.legend()
            fig_querversatz.savefig(querversatz_buf, format='png')
            plt.close(fig_querversatz)
            

            if lateral_1p07 < 1.83:
                warnings.append(f"WARNING: Querversatz 1.07s after start of SWD is smaller than 1,83 meter for signal pair {signal_name_pair}")
                break
            
    if not signal_name_querversatz:
        plt.close(fig_querversatz)
        progress_listbox.insert(tk.END, f"WARNING: Querversatz not available for file: {filename}")
        warnings.append("WARNING: Querversatz not available for this file")
    else:
        plot_buf.append(querversatz_buf)
    
    # Check vx in km/h
    vx_begin_kmh = mf4_data_reduced[signal_name_vx][0] * 3.6
    vx_within_80_pm_2_kmh = (78 <= vx_begin_kmh <= 82)

    # Check Schwimmwinkel
    signal_names_schwimmwinkel = ["DcrInEgoM_beta_Act", "Ins_beta_Act"]
    units_schwimmwinkel = ["rad", "rad"]
    signal_name_schwimmwinkel = ""
    unit_schwimmwinkel = ""

    for signal_name, unit in zip(signal_names_schwimmwinkel, units_schwimmwinkel):
        if signal_name in mf4_data_reduced.columns:
            signal_name_schwimmwinkel += signal_name + ","
            unit_schwimmwinkel += unit + ","
            schwimmwinkel = mf4_data_reduced[signal_name]
            if unit == "rad":
                schwimmwinkel = np.rad2deg(schwimmwinkel)
            if np.any(np.abs(schwimmwinkel) > 15):
                warnings.append(f"WARNING: Schwimmwinkel is bigger than 15 degrees for signal {signal_name}, maximum absolute value: {np.max(abs(schwimmwinkel))} degrees")
    if not signal_name_schwimmwinkel:
        progress_listbox.insert(tk.END, f"WARNING: Schwimmwinkel not available for file: {filename}")
        warnings.append("WARNING: Schwimmwinkel not available for this file")

    # Criterium for test passed
    test_passed = within_35_percent and within_20_percent and vx_within_80_pm_2_kmh

    # Create a dictionary to store signals used
    signal_info = {
        "psid": {"name": signal_name_psi_d, "unit": unit_psi_d},
        "lenkwinkel": {"name": signal_name_lenkwinkel, "unit": unit_lenkwinkel},
        "vx": {"name": signal_name_vx, "unit": unit_vx},  # still included for completeness
        "schwimmwinkel": {"name": signal_name_schwimmwinkel, "unit": unit_schwimmwinkel},
        "querversatz": {"name": signal_name_querversatz, "unit": unit_querversatz}
    }
    
    progress_listbox.insert(tk.END, f"Processed file: {filename}")

    return {
        "filename": filename,
        "plot": plot_buf,
        "psid_peak": psid_peak,
        "signal_info": signal_info,
        "within_35_percent": within_35_percent,
        "within_20_percent": within_20_percent,
        "vx_begin_kmh": vx_begin_kmh,
        "warning": warnings,
        "test_passed": test_passed
    }

def process_mf4_manipulated_signals_data(mf4_data, filename, progress_listbox):
    plot_buf = []
    manipulation_identified = False

    # Function to process the manipulated signals data
    if mf4_data is None or mf4_data.empty:
        progress_listbox.insert(tk.END, f"Manipulated signals data is empty or invalid for file: {filename}")
        return {
            "filename": filename,
            "plot": None,
        }

    # Check whether rDyn is manipulated -------------------------------------------------------------------
    fig_rDyn, ax_rDyn = plt.subplots(figsize=(10, 6))
    rDyn_is_manipulated = False

    signal_pairs = [
        ('EgoWhlRtab_rFL_Whl', 'EgoWhlRtab_rFL_WhlLut'),
        ('EgoWhlRtab_rFR_Whl', 'EgoWhlRtab_rFR_WhlLut'),
        ('EgoWhlRtab_rRL_Whl', 'EgoWhlRtab_rRL_WhlLut'),
        ('EgoWhlRtab_rRR_Whl', 'EgoWhlRtab_rRR_WhlLut')
    ]
    for actual, lookup in signal_pairs:
        if actual in mf4_data.columns and lookup in mf4_data.columns:
            if any(abs(mf4_data[actual] - mf4_data[lookup]) >= 0.03):
                ax_rDyn.plot(mf4_data['time'], mf4_data[actual], label=actual)
                ax_rDyn.plot(mf4_data['time'], mf4_data[lookup], label=lookup)
                rDyn_is_manipulated = True

    if rDyn_is_manipulated:
        ax_rDyn.set_xlabel("Time (s)")
        ax_rDyn.set_ylabel("rDyn (m)")
        ax_rDyn.legend()
        ax_rDyn.grid()
        rDyn_buf = io.BytesIO()
        fig_rDyn.savefig(rDyn_buf, format='png')
        plt.close(fig_rDyn)
        plot_buf.append(rDyn_buf)

        manipulation_identified = True
    else:
        plt.close(fig_rDyn)
    # Check whether rDyn is manipulated ----------------------------------------------------------------------
    
    # Check whether vx or vy is manipulated ------------------------------------------------------------------
    fig_v, ax_v = plt.subplots(figsize=(10, 6))
    v_is_manipulated = False

    nonmanipulated_vx = None
    if "EgoIn5ms_Ins_vx_Act" in mf4_data.columns:
        nonmanipulated_vx = "EgoIn5ms_Ins_vx_Act"
    elif "MobsObs_vx_Act" in mf4_data.columns:
        nonmanipulated_vx = "MobsObs_vx_Act"
    if nonmanipulated_vx is not None and "DcrInEgoM_vx_Act" in mf4_data.columns:
        difference = np.abs(mf4_data[nonmanipulated_vx] - mf4_data["DcrInEgoM_vx_Act"])
        index = np.where(np.diff(difference) >= 0.9)[0]
        first_index = index[0] if index.size > 0 else None
        if first_index is not None:
            ax_v.plot(mf4_data['time'][first_index - 10:first_index + 10], mf4_data[nonmanipulated_vx][first_index - 10:first_index + 10], label=nonmanipulated_vx)
            ax_v.plot(mf4_data['time'][first_index - 10:first_index + 10], mf4_data['DcrInEgoM_vx_Act'][first_index - 10:first_index + 10], label='DcrInEgoM_vx_Act')
            v_is_manipulated = True
    
    nonmanipulated_vy = None
    if "EgoIn5ms_Ins_vy_Act" in mf4_data.columns:
        nonmanipulated_vy = "EgoIn5ms_Ins_vy_Act"
    elif "MobsObs_vy_Act" in mf4_data.columns:
        nonmanipulated_vy = "MobsObs_vy_Act"
    if nonmanipulated_vy is not None and "DcrInEgoM_vy_Act" in mf4_data.columns:
        difference = np.abs(mf4_data[nonmanipulated_vy] - mf4_data["DcrInEgoM_vy_Act"])
        index = np.where(np.diff(difference) >= 0.5)[0]
        first_index = index[0] if index.size > 0 else None
        if first_index is not None:
            ax_v.plot(mf4_data['time'][first_index - 10:first_index + 10], mf4_data[nonmanipulated_vy][first_index - 10:first_index + 10], label=nonmanipulated_vy)
            ax_v.plot(mf4_data['time'][first_index - 10:first_index + 10], mf4_data['DcrInEgoM_vy_Act'][first_index - 10:first_index + 10], label='DcrInEgoM_vy_Act')
            v_is_manipulated = True

    if v_is_manipulated:
        ax_v.set_xlabel("Time (s)")
        ax_v.set_ylabel("Velocity (m/s)")
        ax_v.legend()
        ax_v.grid()
        v_buf = io.BytesIO()
        fig_v.savefig(v_buf, format='png')
        plt.close(fig_rDyn)
        plot_buf.append(v_buf)

        manipulation_identified = True
    else:
        plt.close(fig_v)
    # Check whether vx or vy is manipulated ------------------------------------------------------------------
    
    # Check whether steering angle is manipulated -----------------------------------------------------------------
    fig_stea, ax_stea = plt.subplots(figsize=(10, 6))
    stea_is_manipulated = False

    if "EgoSteaProc_agFA_Whl" in mf4_data.columns and "EgoStea_agFA_WhlRaw" in mf4_data.columns:
        if any(abs(mf4_data["EgoSteaProc_agFA_Whl"] - mf4_data["EgoStea_agFA_WhlRaw"]) >= 0.005):
            ax_stea.plot(mf4_data['time'], mf4_data['EgoSteaProc_agFA_Whl'], label='EgoSteaProc_agFA_Whl')
            ax_stea.plot(mf4_data['time'], mf4_data['EgoStea_agFA_WhlRaw'], label='EgoStea_agFA_WhlRaw')
            stea_is_manipulated = True
        
    # if "EgoStea_agFA_WhlOfs" in mf4_data.columns:
    #     if any(abs(np.diff(mf4_data["EgoStea_agFA_WhlOfs"])) >= 0.001):
    #         ax_stea.plot(mf4_data['time'], mf4_data['EgoStea_agFA_WhlOfs'], label='EgoStea_agFA_WhlOfs')
    #         stea_is_manipulated = True
    # if "EgoStea_agFA_WhlOfsLt" in mf4_data.columns:
    #     if any(abs(np.diff(mf4_data["EgoStea_agFA_WhlOfsLt"])) >= 0.001):
    #         ax_stea.plot(mf4_data['time'], mf4_data['EgoStea_agFA_WhlOfsLt'], label='EgoStea_agFA_WhlOfsLt')
    #         stea_is_manipulated = True
    # if "EgoStea_agFA_WhlOfsPni" in mf4_data.columns:
    #     if any(abs(np.diff(mf4_data["EgoStea_agFA_WhlOfsPni"])) >= 0.1):
    #         ax_ofspn.plot(mf4_data['time'], mf4_data['EgoStea_agFA_WhlOfsPni'], label='EgoStea_agFA_WhlOfsPni')
    #         stea_is_manipulated = True
    # if "EgoStea_agFA_WhlOfsPsid" in mf4_data.columns:
    #     if any(abs(np.diff(mf4_data["EgoStea_agFA_WhlOfsPsid"])) >= 0.001):
    #         ax_ofspn.plot(mf4_data['time'], mf4_data['EgoStea_agFA_WhlOfsPsid'], label='EgoStea_agFA_WhlOfsPsid')
    #         stea_is_manipulated = True

    if stea_is_manipulated:
        ax_stea.set_xlabel("Time (s)")
        ax_stea.set_ylabel("Steering angle (rad)")
        ax_stea.legend()
        ax_stea.grid()
        
        stea_buf = io.BytesIO()
        fig_stea.savefig(stea_buf, format='png')
        plt.close(fig_stea)
        plot_buf.append(stea_buf)

        manipulation_identified = True
    else:
        plt.close(fig_stea)
    # Check whether steering wheel is manipulated ------------------------------------------------------------------

    # Check whether alpha (slip angle) is manipulated ---------------------------------------------------------------
    fig_alpha, ax_alpha = plt.subplots(2, 2, figsize=(10, 12))
    alpha_is_manipulated = False

    # Plot for Front Left (FL)
    if "SnsMobsAdj_alphaFL_Act" in mf4_data.columns and "EgoOut_alpha_Act_rc.EgoOut_alphaFL_Act" in mf4_data.columns:
        if any(abs(mf4_data["SnsMobsAdj_alphaFL_Act"] - mf4_data["EgoOut_alpha_Act_rc.EgoOut_alphaFL_Act"]) >= 0.05):
            ax_alpha[0, 0].plot(mf4_data['time'], mf4_data['SnsMobsAdj_alphaFL_Act'], label='SnsMobsAdj_alphaFL_Act')
            ax_alpha[0, 0].plot(mf4_data['time'], mf4_data['EgoOut_alpha_Act_rc.EgoOut_alphaFL_Act'], label='EgoOut_alpha_Act_rc.EgoOut_alphaFL_Act')
            ax_alpha[0, 0].set_title("Front Left (FL)")
            ax_alpha[0, 0].set_xlabel("Time (s)")
            ax_alpha[0, 0].set_ylabel("Alpha (rad)")
            ax_alpha[0, 0].legend()
            ax_alpha[0, 0].grid()
            alpha_is_manipulated = True

    # Plot for Front Right (FR)
    if "SnsMobsAdj_alphaFR_Act" in mf4_data.columns and "EgoOut_alpha_Act_rc.EgoOut_alphaFR_Act" in mf4_data.columns:
        if any(abs(mf4_data["SnsMobsAdj_alphaFR_Act"] - mf4_data["EgoOut_alpha_Act_rc.EgoOut_alphaFR_Act"]) >= 0.05):
            ax_alpha[0, 1].plot(mf4_data['time'], mf4_data['SnsMobsAdj_alphaFR_Act'], label='SnsMobsAdj_alphaFR_Act')
            ax_alpha[0, 1].plot(mf4_data['time'], mf4_data['EgoOut_alpha_Act_rc.EgoOut_alphaFR_Act'], label='EgoOut_alpha_Act_rc.EgoOut_alphaFR_Act')
            ax_alpha[0, 1].set_title("Front Right (FR)")
            ax_alpha[0, 1].set_xlabel("Time (s)")
            ax_alpha[0, 1].set_ylabel("Alpha (rad)")
            ax_alpha[0, 1].legend()
            ax_alpha[0, 1].grid()
            alpha_is_manipulated = True

    # Plot for Rear Left (RL)
    if "SnsMobsAdj_alphaRL_Act" in mf4_data.columns and "EgoOut_alpha_Act_rc.EgoOut_alphaRL_Act" in mf4_data.columns:
        if any(abs(mf4_data["SnsMobsAdj_alphaRL_Act"] - mf4_data["EgoOut_alpha_Act_rc.EgoOut_alphaRL_Act"]) >= 0.05):
            ax_alpha[1, 0].plot(mf4_data['time'], mf4_data['SnsMobsAdj_alphaRL_Act'], label='SnsMobsAdj_alphaRL_Act')
            ax_alpha[1, 0].plot(mf4_data['time'], mf4_data['EgoOut_alpha_Act_rc.EgoOut_alphaRL_Act'], label='EgoOut_alpha_Act_rc.EgoOut_alphaRL_Act')
            ax_alpha[1, 0].set_title("Rear Left (RL)")
            ax_alpha[1, 0].set_xlabel("Time (s)")
            ax_alpha[1, 0].set_ylabel("Alpha (rad)")
            ax_alpha[1, 0].legend()
            ax_alpha[1, 0].grid()
            alpha_is_manipulated = True

    # Plot for Rear Right (RR)
    if "SnsMobsAdj_alphaRR_Act" in mf4_data.columns and "EgoOut_alpha_Act_rc.EgoOut_alphaRR_Act" in mf4_data.columns:
        if any(abs(mf4_data["SnsMobsAdj_alphaRR_Act"] - mf4_data["EgoOut_alpha_Act_rc.EgoOut_alphaRR_Act"]) >= 0.05):
            ax_alpha[1, 1].plot(mf4_data['time'], mf4_data['SnsMobsAdj_alphaRR_Act'], label='SnsMobsAdj_alphaRR_Act')
            ax_alpha[1, 1].plot(mf4_data['time'], mf4_data['EgoOut_alpha_Act_rc.EgoOut_alphaRR_Act'], label='EgoOut_alpha_Act_rc.EgoOut_alphaRR_Act')
            ax_alpha[1, 1].set_title("Rear Right (RR)")
            ax_alpha[1, 1].set_xlabel("Time (s)")
            ax_alpha[1, 1].set_ylabel("Alpha (rad)")
            ax_alpha[1, 1].legend()
            ax_alpha[1, 1].grid()
            alpha_is_manipulated = True

    if alpha_is_manipulated:
        plt.tight_layout()
        alpha_buf = io.BytesIO()
        fig_alpha.savefig(alpha_buf, format='png')
        plt.close(fig_alpha)
        plot_buf.append(alpha_buf)

        manipulation_identified = True
    else:
        plt.close(fig_alpha)
    # Check whether alpha (slip angle) is manipulated ---------------------------------------------------------------

    # Check whether psid is manipulated ---------------------------------------------------------------
    # FRAGE: psid is manipiliert mit einem multiplier von z.B. 0.07, nicht addition
    fig_psid, ax_psid = plt.subplots(figsize=(10, 6))
    psid_is_manipulated = False
    if "EgoIn5ms_psid_Raw1" in mf4_data.columns and "EgoInWr_psid_Raw1" in mf4_data.columns:
        if(any(abs(mf4_data["EgoIn5ms_psid_Raw1"] - mf4_data["EgoInWr_psid_Raw1"]) >= 0.0005)):
            ax_psid.plot(mf4_data['time'], mf4_data['EgoIn5ms_psid_Raw1'], label='EgoIn5ms_psid_Raw1')
            ax_psid.plot(mf4_data['time'], mf4_data['EgoInWr_psid_Raw1'], label='EgoInWr_psid_Raw1')
            psid_is_manipulated = True

    if psid_is_manipulated:
        ax_psid.set_xlabel("Time (s)")
        ax_psid.set_ylabel("Psid (rad/s)")
        ax_psid.legend()
        ax_psid.grid()
        psid_buf = io.BytesIO()
        fig_psid.savefig(psid_buf, format='png')
        plt.close(fig_psid)
        plot_buf.append(psid_buf)

        manipulation_identified = True
    else:
        plt.close(fig_psid)
    # Check whether psid is manipulated ---------------------------------------------------------------

    # Check whether ax is manipulated ---------------------------------------------------------------
    # FRAGE: ax is manipiliert mit einem multiplier von z.B. 0.07, nicht addition
    fig_ax, ax_ax = plt.subplots(figsize=(10, 6))
    ax_is_manipulated = False
    if "EgoIn5ms_ax_Raw1" in mf4_data.columns and "EgoInWr_ax_Raw1" in mf4_data.columns:
        if(any(abs(mf4_data["EgoIn5ms_ax_Raw1"] - mf4_data["EgoInWr_ax_Raw1"]) >= 0.05)):
            ax_ax.plot(mf4_data['time'], mf4_data['EgoIn5ms_ax_Raw1'], label='EgoIn5ms_ax_Raw1')
            ax_ax.plot(mf4_data['time'], mf4_data['EgoInWr_ax_Raw1'], label='EgoInWr_ax_Raw1')
            ax_is_manipulated = True

    if ax_is_manipulated:
        ax_ax.set_xlabel("Time (s)")
        ax_ax.set_ylabel("a_x (m/s^2)")
        ax_ax.legend()
        ax_ax.grid()
        ax_buf = io.BytesIO()
        fig_ax.savefig(ax_buf, format='png')
        plt.close(fig_ax)
        plot_buf.append(ax_buf)

        manipulation_identified = True
    else:
        plt.close(fig_ax)
    # Check whether ax is manipulated ---------------------------------------------------------------

    # Check whether ay is manipulated ---------------------------------------------------------------
    fig_ay, ax_ay = plt.subplots(figsize=(10, 6))
    ay_is_manipulated = False
    if "EgoIn5ms_ay_Raw1" in mf4_data.columns and "EgoInWr_ay_Raw1" in mf4_data.columns:
        if(any(abs(np.diff(mf4_data["EgoIn5ms_ay_Raw1"] - mf4_data["EgoInWr_ay_Raw1"])) >= 0.5)):
            ax_ay.plot(mf4_data['time'], mf4_data['EgoIn5ms_ay_Raw1'], label='EgoIn5ms_ay_Raw1')
            ax_ay.plot(mf4_data['time'], mf4_data['EgoInWr_ay_Raw1'], label='EgoInWr_ay_Raw1')
            ay_is_manipulated = True

    if ay_is_manipulated:
        ax_ay.set_xlabel("Time (s)")
        ax_ay.set_ylabel("a_y (m/s^2)")
        ax_ay.legend()
        ax_ay.grid()
        ay_buf = io.BytesIO()
        fig_ay.savefig(ay_buf, format='png')
        plt.close(fig_ay)
        plot_buf.append(ay_buf)

        manipulation_identified = True
    else:
        plt.close(fig_ay)
    # Check whether ay is manipulated ---------------------------------------------------------------

    # Check whether mue split is manipulated ---------------------------------------------------------------
    fig_muesp_1, ax_muesp_1 = plt.subplots(figsize=(10, 6))
    muesp1_is_manipulated = False
    if "DmcIndMs_b_IndMs" in mf4_data.columns:
        if any(mf4_data["DmcIndMs_b_IndMs"] != b'false'):
            ax_muesp_1.plot(mf4_data['time'], mf4_data['DmcIndMs_b_IndMs'], label='DmcIndMs_b_IndMs')
            ax_muesp_1.set_xlabel("Time (s)")
            ax_muesp_1.set_ylabel("Signal Value (bool)")
            ax_muesp_1.legend()
            ax_muesp_1.grid()
            muesp1_is_manipulated = True
    if "DmcIndMs_b_IndMsSusp" in mf4_data.columns:
        if any(mf4_data["DmcIndMs_b_IndMsSusp"] != b'false'):
            ax_muesp_1.plot(mf4_data['time'], mf4_data['DmcIndMs_b_IndMsSusp'], label='DmcIndMs_b_IndMsSusp')
            ax_muesp_1.set_xlabel("Time (s)")
            ax_muesp_1.set_ylabel("Signal Value (bool)")
            ax_muesp_1.legend()
            ax_muesp_1.grid()
            muesp1_is_manipulated = True
    if "DmcIndMs_b_IndMsSusp_Set" in mf4_data.columns:
        if any(mf4_data["DmcIndMs_b_IndMsSusp_Set"] != b'false'):
            ax_muesp_1.plot(mf4_data['time'], mf4_data['DmcIndMs_b_IndMsSusp_Set'], label='DmcIndMs_b_IndMsSusp_Set')
            ax_muesp_1.set_xlabel("Time (s)")
            ax_muesp_1.set_ylabel("Signal Value (bool)")
            ax_muesp_1.legend()
            ax_muesp_1.grid()
            muesp1_is_manipulated = True

    if muesp1_is_manipulated:
        muesp1_buf = io.BytesIO()
        fig_muesp_1.savefig(muesp1_buf, format='png')
        plt.close(fig_muesp_1)
        plot_buf.append(muesp1_buf)

        manipulation_identified = True
    else:
        plt.close(fig_muesp_1)

    fig_muesp_2, ax_muesp_2 = plt.subplots(figsize=(10, 6))
    muesp2_is_manipulated = False
    if "DmcIndMs_fac_IndMs" in mf4_data.columns:
        if(any(abs(np.diff(mf4_data["DmcIndMs_fac_IndMs"]))) > 0):
            ax_muesp_2.plot(mf4_data['time'], mf4_data['DmcIndMs_fac_IndMs'], label='DmcIndMs_fac_IndMs')
            ax_muesp_2.set_xlabel("Time (s)")
            ax_muesp_2.set_ylabel("Signal Value (-)")
            ax_muesp_2.legend()
            ax_muesp_2.grid()
            muesp2_is_manipulated = True
    if "DmcIndMs_fac_IndMsSusp" in mf4_data.columns:
        if(any(abs(np.diff(mf4_data["DmcIndMs_fac_IndMsSusp"]))) > 0):
            ax_muesp_2.plot(mf4_data['time'], mf4_data['DmcIndMs_fac_IndMsSusp'], label='DmcIndMs_fac_IndMsSusp')
            ax_muesp_2.set_xlabel("Time (s)")
            ax_muesp_2.set_ylabel("Signal Value (-)")
            ax_muesp_2.legend()
            ax_muesp_2.grid()
            muesp2_is_manipulated = True
    if "DmcIndMs_cw_IndMs" in mf4_data.columns:
        if(any(abs(np.diff(mf4_data["DmcIndMs_cw_IndMs"]))) > 0):
            ax_muesp_2.plot(mf4_data['time'], mf4_data['DmcIndMs_cw_IndMs'], label='DmcIndMs_cw_IndMs')
            ax_muesp_2.set_xlabel("Time (s)")
            ax_muesp_2.set_ylabel("Signal Value (-)")
            ax_muesp_2.legend()
            ax_muesp_2.grid()
            muesp2_is_manipulated = True
    if "DmcIndMs_cw_IndMs_Side" in mf4_data.columns:
        if(any(abs(np.diff(mf4_data["DmcIndMs_cw_IndMs_Side"]))) > 0):
            ax_muesp_2.plot(mf4_data['time'], mf4_data['DmcIndMs_cw_IndMs_Side'], label='DmcIndMs_cw_IndMs_Side')
            ax_muesp_2.set_xlabel("Time (s)")
            ax_muesp_2.set_ylabel("Signal Value (-)")
            ax_muesp_2.legend()
            ax_muesp_2.grid()
            muesp2_is_manipulated = True

    if muesp2_is_manipulated:
        muesp2_buf = io.BytesIO()
        fig_muesp_2.savefig(muesp2_buf, format='png')
        plt.close(fig_muesp_2)
        plot_buf.append(muesp2_buf)

        manipulation_identified = True
    else:
        plt.close(fig_muesp_2)
    # Check whether mue split is manipulated ---------------------------------------------------------------

    # Check whether Reibwert (Glasbaustein) is manipulated -------------------------------------------------
    # FRAGE: Wenn alle werte kleiner 0.6 oder wenn mindestens nur ein Wert? Betrag oder kein Betrag?
    fig_rw, ax_rw = plt.subplots(figsize=(10, 6))
    rw_is_manipulated = False
    if "SnsRc_mue_ActHi" in mf4_data.columns:
        if(all(mf4_data["SnsRc_mue_ActHi"] < 0.6)):
            ax_rw.plot(mf4_data['time'], mf4_data['SnsRc_mue_ActHi'], label='SnsRc_mue_ActHi')
            rw_is_manipulated = True

    if rw_is_manipulated:
        ax_rw.set_xlabel("Time (s)")
        ax_rw.set_ylabel("Signal Value (-)")
        ax_rw.legend()
        ax_rw.grid()
        rw_buf = io.BytesIO()
        fig_rw.savefig(rw_buf, format='png')
        plt.close(fig_rw)
        plot_buf.append(rw_buf)

        manipulation_identified = True
    else:
        plt.close(fig_rw)
    # Check whether Reibwert (Glasbaustein) is manipulated -------------------------------------------------

    # Check whether ay no grav is manipulated -------------------------------------------------
    fig_ay_no_grav, ax_ay_no_grav = plt.subplots(figsize=(10, 6))
    ay_no_grav_is_manipulated = False
    if "DcrInEgoM_ay_ActNoGrav" in mf4_data.columns and "DcrInEgoM_ay_Act" in mf4_data.columns:
        if all(mf4_data["DcrInEgoM_ay_ActNoGrav"] == 0):
            ax_ay_no_grav.plot(mf4_data['time'], mf4_data['DcrInEgoM_ay_ActNoGrav'], label='DcrInEgoM_ay_ActNoGrav')
            ay_no_grav_is_manipulated = True
    if ay_no_grav_is_manipulated:
        ax_ay_no_grav.set_xlabel("Time (s)")
        ax_ay_no_grav.set_ylabel("a_y_no_grav (m/s^2)")
        ax_ay_no_grav.legend()
        ax_ay_no_grav.grid()
        ay_no_grav_buf = io.BytesIO()
        fig_ay_no_grav.savefig(ay_no_grav_buf, format='png')
        plt.close(fig_ay_no_grav)
        plot_buf.append(ay_no_grav_buf)
        
        manipulation_identified = True
    else:
        plt.close(fig_ay_no_grav)
    # Check whether ay no grav is manipulated -------------------------------------------------

    # Check whether TEE is manipulated -------------------------------------------------

    # Define FA and RA signal groups
    fa_signals = [
        "DcrTql_tqwFA_EmMaxDyn",
        "DcrTql_tqwFA_EmMaxStat",
        "DcrTql_tqwFA_EmMinDyn",
        "DcrTql_tqwFA_EmMinStat"
    ]
    ra_signals = [
        "DcrTql_tqwRA_EmMaxDyn",
        "DcrTql_tqwRA_EmMaxStat",
        "DcrTql_tqwRA_EmMinDyn",
        "DcrTql_tqwRA_EmMinStat"
    ]

    # Plot FA signals if all are always zero
    fa_all_zero = all(
        signal in mf4_data.columns and np.all(mf4_data[signal] == 0)
        for signal in fa_signals
    )
    if fa_all_zero:
        fig_tee_fa, ax_tee_fa = plt.subplots(figsize=(10, 6))
        for signal in fa_signals:
            if signal in mf4_data.columns:
                ax_tee_fa.plot(mf4_data['time'], mf4_data[signal], label=signal)
        ax_tee_fa.set_title("TEE FA signals (all zero)")
        ax_tee_fa.set_xlabel("Time (s)")
        ax_tee_fa.set_ylabel("Moment (Nm)")
        ax_tee_fa.legend()
        ax_tee_fa.grid()
        tee_fa_buf = io.BytesIO()
        fig_tee_fa.savefig(tee_fa_buf, format='png')
        plt.close(fig_tee_fa)
        plot_buf.append(tee_fa_buf)

        manipulation_identified = True

    # Plot RA signals if all are always zero
    ra_all_zero = all(
        signal in mf4_data.columns and np.all(mf4_data[signal] == 0)
        for signal in ra_signals
    )
    if ra_all_zero:
        fig_tee_ra, ax_tee_ra = plt.subplots(figsize=(10, 6))
        for signal in ra_signals:
            if signal in mf4_data.columns:
                ax_tee_ra.plot(mf4_data['time'], mf4_data[signal], label=signal)
        ax_tee_ra.set_title("TEE RA signals (all zero)")
        ax_tee_ra.set_xlabel("Time (s)")
        ax_tee_ra.set_ylabel("Moment (Nm)")
        ax_tee_ra.legend()
        ax_tee_ra.grid()
        tee_ra_buf = io.BytesIO()
        fig_tee_ra.savefig(tee_ra_buf, format='png')
        plt.close(fig_tee_ra)
        plot_buf.append(tee_ra_buf)

        manipulation_identified = True

    # Check whether TEE is manipulated -------------------------------------------------

    # # Check whether vch is manipulated -------------------------------------------------
    # fig_vch, ax_vch = plt.subplots(figsize=(10, 6))
    # vch_is_manipulated = False
    # if "EgoOutWr_rc_v_Ch.EgoOut_v_Ch" in mf4_data.columns and "DcrInEgoE_vch_Veh" in mf4_data.columns:
    #     if(any(abs(np.diff(mf4_data["EgoOutWr_rc_v_Ch.EgoOut_v_Ch"] - mf4_data["DcrInEgoE_vch_Veh"])) >= 0.01)):
    #         ax_vch.plot(mf4_data['time'], mf4_data['EgoOutWr_rc_v_Ch.EgoOut_v_Ch'], label='EgoOutWr_rc_v_Ch.EgoOut_v_Ch')
    #         ax_vch.plot(mf4_data['time'], mf4_data['DcrInEgoE_vch_Veh'], label='DcrInEgoE_vch_Veh')
    #         vch_is_manipulated = True
    # if vch_is_manipulated:
    #     ax_vch.set_xlabel("Time (s)")
    #     ax_vch.set_ylabel("Signal Value (-)")
    #     ax_vch.legend()
    #     ax_vch.grid()
    #     vch_buf = io.BytesIO()
    #     fig_vch.savefig(vch_buf, format='png')
    #     plt.close(fig_vch)
    #     plot_buf.append(vch_buf)
    # CLOSE PLOT AND dont forget manipulation_identified
    # # Check whether vch is manipulated -------------------------------------------------

    return {
        "filename": filename,
        "plot": plot_buf,
        "manipulation_identified": manipulation_identified,
    }

def create_word_document(analysis_data_list, manipulated_signals_data_list, output_filename):
    # Function to create a Word document with analysis results
    doc = Document()
    doc.add_heading('MF4 Data Analysis', 0)

    for (analysis_data, manipulated_signals_data) in zip(analysis_data_list, manipulated_signals_data_list):

        if not manipulated_signals_data["manipulation_identified"]:
            heading = doc.add_heading(level=1)
            run = heading.add_run(f'Unknown manipulation in File: {os.path.basename(analysis_data["filename"])}')
            run.bold = True  # Make text bold
            run.font.color.rgb = RGBColor(255, 0, 0)  # Make text red
        else:
            doc.add_heading(f'File: {os.path.basename(analysis_data["filename"])}', level=1)

        if analysis_data["warning"] is not None:
            for warning in analysis_data["warning"]:
                doc.add_paragraph(warning, style='ListBullet')

        if analysis_data["plot"] is not None:
            for plot_buf in analysis_data["plot"]:
                plot_buf.seek(0)  # Reset the file pointer to the beginning
                doc.add_picture(plot_buf, width=Inches(6))
                doc.add_paragraph()
        
        if len(manipulated_signals_data["plot"]) > 0:
            doc.add_paragraph('Manipulated Signals:')
            for plot_buf in manipulated_signals_data["plot"]:
                plot_buf.seek(0)
                doc.add_picture(plot_buf, width=Inches(6))
                doc.add_paragraph()

        if analysis_data["signal_info"] is not None:
            doc.add_paragraph(f'Psid Signal Used: {analysis_data["signal_info"]["psid"]["name"]} ({analysis_data["signal_info"]["psid"]["unit"]})')

            doc.add_paragraph(f'Lenkwinkel Signal Used: {analysis_data["signal_info"]["lenkwinkel"]["name"]} ({analysis_data["signal_info"]["lenkwinkel"]["unit"]})')

            doc.add_paragraph(f'Vx Signal Used: {analysis_data["signal_info"]["vx"]["name"]} ({analysis_data["signal_info"]["vx"]["unit"]})')

            doc.add_paragraph(f'Schwimmwinkel Signal Used: {analysis_data["signal_info"]["schwimmwinkel"]["name"]} ({analysis_data["signal_info"]["schwimmwinkel"]["unit"]})')

            doc.add_paragraph(f'Querversatz Signal Used: {analysis_data["signal_info"]["querversatz"]["name"]} ({analysis_data["signal_info"]["querversatz"]["unit"]})')

            doc.add_paragraph(f'Psid Peak: {analysis_data["psid_peak"]:.3f} {analysis_data["signal_info"]["psid"]["unit"]}')

        if analysis_data["within_35_percent"] is not None:
            doc.add_paragraph(f'Psid at T_0+1 is within 35% range: {analysis_data["within_35_percent"]}')

        if analysis_data["within_20_percent"] is not None:
            doc.add_paragraph(f'Psid at T_0+1.75 is within 20% range: {analysis_data["within_20_percent"]}')

        if analysis_data["vx_begin_kmh"] is not None:
            doc.add_paragraph(f'Vx at in the beginning of SWD: {analysis_data["vx_begin_kmh"]:.3f} km/h')

        if analysis_data["test_passed"] is not None:
            doc.add_paragraph(f'Test Passed: {analysis_data["test_passed"]}')

    doc.save(output_filename)
    print("Output saved to", os.path.abspath(output_filename))

def select_folder_and_process():
    # Function to select a folder and process MF4 files
    root = tk.Tk()
    root.title("MF4 File Processor")

    # Create a listbox to show the processing progress
    progress_listbox = Listbox(root, width=80, height=20)
    progress_listbox.pack()

    folder_path = filedialog.askdirectory(title="Select Folder with MF4 Files")
    if not folder_path:
        messagebox.showerror("Error", "No folder selected.")
        root.destroy()
        return

    analysis_data_list = []
    manipulated_signals_data_list = []
    for file in os.listdir(folder_path):
        if file.endswith(".mf4"):
            mf4_file = os.path.join(folder_path, file)
            progress_listbox.insert(tk.END, f"Processing file: {mf4_file}")
            root.update_idletasks()  # Update the GUI to show progress

            mf4_analysis_data, mf4_manipulated_signals_data = extract_data(mf4_file, convert=True)
            analysis_data = process_mf4_analysis_data(mf4_analysis_data, mf4_file, progress_listbox)
            manipulated_signals_data = process_mf4_manipulated_signals_data(mf4_manipulated_signals_data, mf4_file, progress_listbox)
            if analysis_data is not None:
                analysis_data_list.append(analysis_data)
            if manipulated_signals_data is not None:
                manipulated_signals_data_list.append(manipulated_signals_data)

    if len(analysis_data_list) > 0 and len(manipulated_signals_data_list) > 0:
        output_filename = os.path.join(folder_path, "SWD_Analysis_Report.docx")
        for manipulated_signals_data in manipulated_signals_data_list:
            if not manipulated_signals_data["manipulation_identified"]:
                output_filename = os.path.join(folder_path, "SWD_Analysis_Report_invalidManip.docx")
                break

        create_word_document(analysis_data_list, manipulated_signals_data_list,output_filename)
        progress_listbox.insert(tk.END, f"Analysis complete. Report saved to:\n{os.path.abspath(output_filename)}")
        messagebox.showinfo("Success", f"Analysis complete. Report saved to:\n{os.path.abspath(output_filename)}")
    else:
        progress_listbox.insert(tk.END, "No valid MF4 data processed.")
        messagebox.showinfo("Info", "No valid MF4 data processed.")

    root.mainloop()

if __name__ == "__main__":
    select_folder_and_process()